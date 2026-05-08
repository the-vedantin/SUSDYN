"""
vahan/report_gen.py  —  Vehicle Dynamics Report Generator
==========================================================
Produces a .docx that opens and edits cleanly in Google Docs.

Call::

    from vahan.report_gen import generate_report
    generate_report(output_path, data)

where ``data`` is assembled by ``MainWindow._collect_report_data()``.

No Qt imports here — pure NumPy / Matplotlib / python-docx.
"""

from __future__ import annotations

import io
import datetime
import math
from typing import Any

import numpy as np

import matplotlib
matplotlib.use('Agg')                        # headless rendering
import matplotlib.pyplot as plt
import matplotlib.ticker as mticker

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── print-friendly colours ────────────────────────────────────────────────────
_FL  = '#1565C0'   # dark blue
_FR  = '#42A5F5'   # sky blue
_RL  = '#C62828'   # dark red
_RR  = '#EF5350'   # medium red
_ELF = '#1565C0'   # elastic LT front
_ELR = '#C62828'   # elastic LT rear
_GLF = '#F57F17'   # geometric LT front
_GLR = '#6A1B9A'   # geometric LT rear

_CORNER_COLOR = {'FL': _FL, 'FR': _FR, 'RL': _RL, 'RR': _RR}
_CORNER_LS    = {'FL': '-', 'FR': '--', 'RL': '-.', 'RR': ':'}

REPORT_FIG_W = 5.5   # inches, fits letter page with 1-in margins


# =============================================================================
# ── DOCX helpers ─────────────────────────────────────────────────────────────
# =============================================================================

def _set_cell_bg(cell, hex_color: str):
    """Set cell background colour (e.g. '#F5F5F5')."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color.lstrip('#'))
    tcPr.append(shd)


def _set_cell_border(cell, **kwargs):
    """Set borders on a table cell.  kwargs: top/bottom/left/right = {'sz':4,'val':'single'}."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, attrs in kwargs.items():
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   attrs.get('val',   'single'))
        el.set(qn('w:sz'),    str(attrs.get('sz', 4)))
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), attrs.get('color', 'auto'))
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _hr(doc: Document):
    """Thin horizontal rule as a paragraph border."""
    p    = doc.add_paragraph()
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '4')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), 'AAAAAA')
    pBdr.append(bot)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(4)


def _body(doc: Document, text: str, italic=False, bold=False, color=None,
          size_pt=10.5, space_after=4):
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name   = 'Arial'
    run.font.size   = Pt(size_pt)
    run.font.italic = italic
    run.font.bold   = bold
    if color:
        r, g, b = int(color[1:3], 16), int(color[3:5], 16), int(color[5:7], 16)
        run.font.color.rgb = RGBColor(r, g, b)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    return p


def _heading(doc: Document, text: str, level: int):
    """Section heading styled for clean Google Docs rendering."""
    p   = doc.add_heading(level=level)
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(16 if level == 1 else 13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x15, 0x65, 0xC0) if level == 1 else RGBColor(0x1A, 0x1A, 0x2E)
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after  = Pt(4)
    return p


def _analysis_box(doc: Document, text: str):
    """Light-gray callout for auto-generated analysis."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    _set_cell_bg(cell, '#F0F4FF')
    border_attrs = {'val': 'single', 'sz': 4, 'color': '90A4AE'}
    for side in ('top', 'bottom', 'left', 'right'):
        _set_cell_border(cell, **{side: border_attrs})
    p   = cell.paragraphs[0]
    run = p.add_run('Analysis:  ')
    run.font.name   = 'Arial'
    run.font.size   = Pt(10)
    run.font.bold   = True
    run.font.color.rgb = RGBColor(0x15, 0x65, 0xC0)
    run2 = p.add_run(text)
    run2.font.name   = 'Arial'
    run2.font.size   = Pt(10)
    run2.font.italic = True
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    cell.add_paragraph('')   # padding
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def _rationale_box(doc: Document):
    """Editable design rationale placeholder."""
    tbl  = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    _set_cell_bg(cell, '#FAFAFA')
    border_attrs = {'val': 'dashed', 'sz': 4, 'color': 'BDBDBD'}
    for side in ('top', 'bottom', 'left', 'right'):
        _set_cell_border(cell, **{side: border_attrs})
    p   = cell.paragraphs[0]
    run = p.add_run('Design Rationale:  ')
    run.font.name  = 'Arial'
    run.font.size  = Pt(10)
    run.font.bold  = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    run2 = p.add_run('Replace this text with your engineering justification.')
    run2.font.name  = 'Arial'
    run2.font.size  = Pt(10)
    run2.font.color.rgb = RGBColor(0xBD, 0xBD, 0xBD)
    for _ in range(3):
        cell.add_paragraph('')
    doc.add_paragraph().paragraph_format.space_after = Pt(8)


def _embed_figure(doc: Document, fig, width_in=REPORT_FIG_W, caption: str = ''):
    """Render matplotlib figure to PNG bytes and embed in doc."""
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight',
                facecolor='white', edgecolor='none')
    buf.seek(0)
    plt.close(fig)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(buf, width=Inches(width_in))
    p.paragraph_format.space_after = Pt(2)
    if caption:
        cp = doc.add_paragraph(caption)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.runs[0].font.name   = 'Arial'
        cp.runs[0].font.size   = Pt(9)
        cp.runs[0].font.italic = True
        cp.runs[0].font.color.rgb = RGBColor(0x75, 0x75, 0x75)
        cp.paragraph_format.space_after = Pt(8)


def _param_table(doc: Document, rows: list[tuple[str, str, str]]):
    """rows = [(label, value, unit), ...]  — simple 3-column table."""
    tbl = doc.add_table(rows=len(rows) + 1, cols=3)
    tbl.style = 'Table Grid'
    # header
    for col, hdr in enumerate(('Parameter', 'Value', 'Unit')):
        c = tbl.cell(0, col)
        c.text = hdr
        _set_cell_bg(c, '#1565C0')
        c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        c.paragraphs[0].runs[0].font.bold = True
        c.paragraphs[0].runs[0].font.name = 'Arial'
        c.paragraphs[0].runs[0].font.size = Pt(10)
    for r_idx, (label, value, unit) in enumerate(rows):
        row_cells = tbl.rows[r_idx + 1].cells
        bg = 'FFFFFF' if r_idx % 2 == 0 else 'F5F5F5'
        for ci, txt in enumerate((label, value, unit)):
            row_cells[ci].text = txt
            _set_cell_bg(row_cells[ci], '#' + bg)
            row_cells[ci].paragraphs[0].runs[0].font.name = 'Arial'
            row_cells[ci].paragraphs[0].runs[0].font.size = Pt(10)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


# =============================================================================
# ── MATPLOTLIB FIGURE GENERATORS ─────────────────────────────────────────────
# =============================================================================

def _fig_style():
    return {
        'facecolor': 'white',
        'figsize':   (REPORT_FIG_W, 3.2),
    }


def _ax_style(ax, xlabel='', ylabel='', title=''):
    ax.set_facecolor('#FAFAFA')
    ax.tick_params(labelsize=8)
    ax.grid(True, color='#E0E0E0', linewidth=0.6, linestyle='--')
    ax.spines[['top', 'right']].set_visible(False)
    if xlabel: ax.set_xlabel(xlabel, fontsize=9)
    if ylabel: ax.set_ylabel(ylabel, fontsize=9)
    if title:  ax.set_title(title, fontsize=10, fontweight='bold', pad=4)


def _fig_kinem_corners(x_arr, sweep_results, key, ylabel, title,
                        xlabel='Wheel Travel (mm)', corners=('FL', 'FR', 'RL', 'RR'),
                        show_front_rear_avg=False):
    """Single kinematic metric across chosen corners."""
    fig, ax = plt.subplots(**_fig_style())
    for lbl in corners:
        data = sweep_results.get(lbl, {}).get(key)
        if data is None or not np.any(np.isfinite(data)):
            continue
        ax.plot(x_arr, data, color=_CORNER_COLOR[lbl],
                linestyle=_CORNER_LS[lbl], linewidth=1.6, label=lbl)
    _ax_style(ax, xlabel, ylabel, title)
    ax.legend(fontsize=8, framealpha=0.7, loc='best')
    fig.tight_layout(pad=0.5)
    return fig


def _fig_dyn_multi(x_arr, series_list, xlabel, title):
    """Generic multi-series dynamic figure.
    series_list = [(label, y_array, color, linestyle, ylabel_group), ...]
    """
    fig, ax = plt.subplots(**_fig_style())
    for label, y_arr, color, ls in series_list:
        ax.plot(x_arr, y_arr, color=color, linestyle=ls, linewidth=1.6, label=label)
    _ax_style(ax, xlabel, '', title)
    ax.legend(fontsize=8, framealpha=0.7, loc='best')
    fig.tight_layout(pad=0.5)
    return fig


def _fig_two_panel(x_arr, top_series, bot_series, xlabel,
                    top_ylabel, bot_ylabel, title):
    """Two-row figure sharing X axis."""
    fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(REPORT_FIG_W, 4.4),
                                    sharex=True, facecolor='white')
    for label, y_arr, color, ls in top_series:
        ax1.plot(x_arr, y_arr, color=color, linestyle=ls, linewidth=1.6, label=label)
    for label, y_arr, color, ls in bot_series:
        ax2.plot(x_arr, y_arr, color=color, linestyle=ls, linewidth=1.6, label=label)
    _ax_style(ax1, '',      top_ylabel, title)
    _ax_style(ax2, xlabel,  bot_ylabel, '')
    ax1.legend(fontsize=8, framealpha=0.7, loc='best')
    ax2.legend(fontsize=8, framealpha=0.7, loc='best')
    fig.tight_layout(pad=0.5)
    return fig


# =============================================================================
# ── AUTO-ANALYSIS TEXT ───────────────────────────────────────────────────────
# =============================================================================

def _pct(value, total):
    if total and total != 0:
        return 100 * value / total
    return 0.0


def _analyze_heave_camber(sweep_results, x_arr) -> str:
    """Camber gain / loss across heave range."""
    parts = []
    for axle_label, corners in [('Front', ('FL', 'FR')), ('Rear', ('RL', 'RR'))]:
        gains = []
        for lbl in corners:
            d = sweep_results.get(lbl, {}).get('camber')
            if d is not None and np.any(np.isfinite(d)):
                mid = len(d) // 2
                gain_per_mm = (d[-1] - d[mid]) / max(abs(x_arr[-1] - x_arr[mid]), 1e-6)
                gains.append(gain_per_mm)
        if gains:
            avg = float(np.mean(gains))
            direction = 'gaining negative' if avg < 0 else 'gaining positive'
            parts.append(f'{axle_label}: {avg:+.3f} deg/mm ({direction} camber in bump)')
    return '  |  '.join(parts) if parts else 'Camber data not available.'


def _analyze_heave_rc(sweep_results, x_arr) -> str:
    rc_f = sweep_results.get('FL', {}).get('rc_height')
    rc_r = sweep_results.get('RL', {}).get('rc_height')
    parts = []
    for label, rc in [('Front RC', rc_f), ('Rear RC', rc_r)]:
        if rc is not None and np.any(np.isfinite(rc)):
            mid = len(rc) // 2
            migration = rc[-1] - rc[mid]
            parts.append(f'{label} migrates {migration:+.1f} mm over full heave range')
    return '  |  '.join(parts) if parts else 'RC height data not available.'


def _analyze_cornering(sweep: dict) -> dict:
    g_arr  = np.asarray(sweep.get('lateral_g', [0]))
    roll   = np.asarray(sweep.get('roll_angle_deg', [0]))
    pitch  = np.asarray(sweep.get('pitch_angle_deg', [0]))
    util   = {c: np.asarray(sweep.get(f'utilization_{c}', np.zeros_like(g_arr)))
              for c in ('FL', 'FR', 'RL', 'RR')}
    fz     = {c: np.asarray(sweep.get(f'Fz_{c}', np.zeros_like(g_arr)))
              for c in ('FL', 'FR', 'RL', 'RR')}

    # Roll
    peak_roll = float(np.nanmax(np.abs(roll))) if len(roll) > 0 else 0.0
    roll_target = '2–3 °' if peak_roll > 3 else ('good' if peak_roll >= 1.5 else 'very stiff')
    if peak_roll > 3.0:
        roll_txt = (f'Peak body roll {peak_roll:.1f} ° exceeds typical 2–3 ° target — '
                    f'consider stiffer ARBs or spring rates.')
    elif peak_roll < 1.0:
        roll_txt = (f'Peak body roll {peak_roll:.1f} ° is very low — '
                    f'good chassis stiffness but may reduce driver feel.')
    else:
        roll_txt = f'Peak body roll {peak_roll:.1f} ° — within a typical FSAE 1–3 ° target.'

    # Utilization — find saturation point
    util_max = np.zeros_like(g_arr)
    for u in util.values():
        util_max = np.maximum(util_max, u)
    sat_idx = np.where(util_max >= 0.95)[0]
    if len(sat_idx) > 0:
        sat_g = float(g_arr[sat_idx[0]])
        # which corner saturates first
        first_corner = max(util, key=lambda c: float(util[c][sat_idx[0]]))
        util_txt = (f'First tire approaches saturation (≥ 95 %) at {sat_g:.2f} g — '
                    f'{first_corner} corner leads.')
    else:
        peak_util = float(np.nanmax(util_max)) if len(util_max) > 0 else 0.0
        util_txt  = (f'No tire reaches saturation in sweep range. '
                     f'Peak utilization {peak_util:.0%} — tires not fully loaded.')

    # Front vs rear inside/outside Fz delta at 1g
    idx_1g = int(np.argmin(np.abs(g_arr - 1.0)))
    fz_fo = float(fz['FR'][idx_1g]); fz_fi = float(fz['FL'][idx_1g])
    fz_ro = float(fz['RR'][idx_1g]); fz_ri = float(fz['RL'][idx_1g])
    lt_f_pct  = _pct(fz_fo - fz_fi, fz_fo + fz_fi)
    lt_r_pct  = _pct(fz_ro - fz_ri, fz_ro + fz_ri)
    lt_txt = (f'At 1 g: front lateral load transfer {lt_f_pct:.0f} % '
              f'(outside/inside Fz), rear {lt_r_pct:.0f} %.')

    return {
        'roll':  roll_txt,
        'util':  util_txt,
        'lt':    lt_txt,
        'pitch': (f'Pitch angle at 1 g lateral: {float(pitch[idx_1g]):.2f} °. '
                  f'Pure cornering pitch change reflects CG height + load transfer asymmetry.'),
    }


def _analyze_traj(sweep: dict, label: str) -> str:
    t_arr  = np.asarray(sweep.get('time_s', [0]))
    v_arr  = np.asarray(sweep.get('speed_mph', [0]))
    g_arr  = np.asarray(sweep.get('longitudinal_g', [0]))
    pitch  = np.asarray(sweep.get('pitch_angle_deg', [0]))
    if len(t_arr) < 2:
        return f'{label}: insufficient data.'
    duration = float(t_arr[-1] - t_arr[0])
    v_final  = float(v_arr[-1])
    v_init   = float(v_arr[0])
    g_peak   = float(np.nanmax(np.abs(g_arr)))
    pitch_pk = float(np.nanmax(np.abs(pitch)))
    if label == 'Acceleration':
        return (f'0–{v_final:.0f} mph in {duration:.1f} s. '
                f'Peak achieved traction {g_peak:.2f} g — limited by drivetrain + traction. '
                f'Nose raises {pitch_pk:.2f} ° at peak accel (weight transfer to rear).')
    else:
        delta_v = abs(v_init - v_final)
        return (f'{v_init:.0f} → {v_final:.0f} mph in {duration:.1f} s '
                f'({delta_v:.0f} mph shed). '
                f'Peak decel {g_peak:.2f} g — friction + aero drag. '
                f'Nose dives {pitch_pk:.2f} ° at peak braking (weight transfer to front).')


# =============================================================================
# ── MAIN GENERATOR ───────────────────────────────────────────────────────────
# =============================================================================

def generate_report(output_path: str, data: dict,
                    progress_cb=None) -> None:
    """
    Generate the Vehicle Dynamics Report and save to ``output_path`` (.docx).

    ``data`` keys (assembled by MainWindow._collect_report_data):
        car_params          dict   — from self._car
        veh_params          dict   — VehicleParams field values
        heave_x_mm          ndarray
        heave_results       dict   — {corner: {key: ndarray}}
        roll_x_deg          ndarray
        roll_results        dict   — {corner: {key: ndarray}}
        dyn_cornering       dict   — sweep_lateral_g result
        dyn_accel           dict   — sweep_acceleration_trajectory result
        dyn_brake           dict   — sweep_acceleration_trajectory result (braking)
        view3d_png          bytes  — PNG screenshot of 3D view (or None)
    """
    def _prog(msg, pct):
        if progress_cb:
            progress_cb(msg, pct)

    doc = Document()

    # ── page margins (1 in all sides) ────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.1)
        section.right_margin  = Inches(1.1)

    car   = data.get('car_params', {})
    vp    = data.get('veh_params', {})
    heave_x   = np.asarray(data.get('heave_x_mm', []))
    heave_res  = data.get('heave_results', {})
    roll_x    = np.asarray(data.get('roll_x_deg', []))
    roll_res   = data.get('roll_results', {})
    dyn_corn   = data.get('dyn_cornering', {})
    dyn_accel  = data.get('dyn_accel', {})
    dyn_brake  = data.get('dyn_brake', {})
    view3d_png = data.get('view3d_png')

    _prog('Cover page…', 2)

    # =========================================================================
    # COVER
    # =========================================================================
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Vehicle Dynamics Report')
    run.font.name = 'Arial'
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x15, 0x65, 0xC0)
    p.paragraph_format.space_before = Pt(40)
    p.paragraph_format.space_after  = Pt(6)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r   = sub.add_run(f"Generated by Vahan  ·  {datetime.date.today().strftime('%B %d, %Y')}")
    r.font.name  = 'Arial'
    r.font.size  = Pt(11)
    r.font.color.rgb = RGBColor(0x75, 0x75, 0x75)
    sub.paragraph_format.space_after = Pt(20)

    if view3d_png:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.add_run().add_picture(io.BytesIO(view3d_png), width=Inches(4.5))
        p2.paragraph_format.space_after = Pt(16)

    doc.add_page_break()

    # =========================================================================
    # SECTION 1 — VEHICLE PARAMETERS
    # =========================================================================
    _prog('Vehicle parameters table…', 8)
    _heading(doc, '1 — Vehicle Parameters', 1)
    _body(doc, 'Key geometry and mass properties used for all simulations.', size_pt=10)
    _hr(doc)

    wb    = car.get('wheelbase_mm', 1537.)
    tf    = car.get('track_f_mm',   1222.)
    tr    = car.get('track_r_mm',   1200.)
    cg_y  = car.get('cg_y_mm',      845.)
    cg_z  = car.get('cg_z_mm',      280.)
    mass  = vp.get('total_mass_kg', 300.)
    m_sp  = vp.get('sprung_mass_kg', mass * 0.85)
    m_us  = vp.get('unsprung_mass_kg', mass - m_sp)
    mu    = vp.get('peak_mu', 1.5)
    pwr   = vp.get('power_hp', 75.)
    wr_f  = vp.get('wheel_rate_front_Npm', 14000.)
    wr_r  = vp.get('wheel_rate_rear_Npm',  14000.)
    front_bias = car.get('front_brake_bias_pct', 65.)

    a_m   = cg_y / 1000.       # CG to front axle
    b_m   = (wb - cg_y) / 1000.
    wf_pct = 100. * b_m / (wb / 1000.)
    wr_pct = 100. * a_m / (wb / 1000.)

    rows = [
        ('Wheelbase',               f'{wb:.0f}',      'mm'),
        ('Front Track',             f'{tf:.0f}',      'mm'),
        ('Rear Track',              f'{tr:.0f}',      'mm'),
        ('CG to Front Axle',        f'{cg_y:.1f}',    'mm'),
        ('CG to Rear Axle',         f'{wb - cg_y:.1f}','mm'),
        ('CG Height',               f'{cg_z:.1f}',    'mm'),
        ('Front Weight Distribution',f'{wf_pct:.1f}', '%'),
        ('Rear Weight Distribution', f'{wr_pct:.1f}', '%'),
        ('Total Mass',              f'{mass:.1f}',    'kg'),
        ('Sprung Mass',             f'{m_sp:.1f}',    'kg'),
        ('Unsprung Mass (est.)',     f'{m_us:.1f}',   'kg'),
        ('Engine Power',            f'{pwr:.0f}',     'hp'),
        ('Peak Tire μ',             f'{mu:.2f}',      '—'),
        ('Wheel Rate (Front)',       f'{wr_f:.0f}',    'N/m'),
        ('Wheel Rate (Rear)',        f'{wr_r:.0f}',    'N/m'),
        ('Front Brake Bias',        f'{front_bias:.0f}', '%'),
    ]
    _param_table(doc, rows)
    _rationale_box(doc)

    # =========================================================================
    # SECTION 2 — SUSPENSION KINEMATICS (HEAVE)
    # =========================================================================
    _prog('Kinematics — heave sweep…', 15)
    _heading(doc, '2 — Suspension Kinematics: Heave', 1)
    _body(doc, (
        'Static vertical travel sweep. Shows how each kinematic quantity changes '
        'as the suspension compresses and extends symmetrically. '
        f'Range: {heave_x[0]:.0f} to {heave_x[-1]:.0f} mm.'
        if len(heave_x) else 'Heave data not available.'
    ), size_pt=10)
    _hr(doc)

    if len(heave_x) and heave_res:

        # 2a — Camber
        _heading(doc, '2.1 — Camber vs Wheel Travel', 2)
        fig = _fig_kinem_corners(heave_x, heave_res, 'camber',
                                  'Camber (°)', 'Camber vs Heave')
        _embed_figure(doc, fig, caption='Negative camber = tire leans inward. Target: gain negative camber in bump.')
        txt = _analyze_heave_camber(heave_res, heave_x)
        _analysis_box(doc, txt)
        _rationale_box(doc)

        # 2b — Toe
        _heading(doc, '2.2 — Toe vs Wheel Travel', 2)
        fig = _fig_kinem_corners(heave_x, heave_res, 'toe',
                                  'Toe (°)', 'Toe vs Heave')
        _embed_figure(doc, fig, caption='Positive toe = toe-out. Bump-toe-in (rear) improves stability.')
        fl_toe = heave_res.get('FL', {}).get('toe')
        rl_toe = heave_res.get('RL', {}).get('toe')
        toe_txt = ''
        for axle_lbl, toe_data in [('Front', fl_toe), ('Rear', rl_toe)]:
            if toe_data is not None and np.any(np.isfinite(toe_data)):
                mid = len(toe_data) // 2
                delta = toe_data[-1] - toe_data[mid]
                direction = 'toe-in' if delta < 0 else 'toe-out'
                toe_txt += f'{axle_lbl}: {delta:+.3f} ° bump {direction} over range.  '
        _analysis_box(doc, toe_txt.strip() or 'Toe data not available.')
        _rationale_box(doc)

        # 2c — Roll Centre Height
        _heading(doc, '2.3 — Roll Centre Height vs Wheel Travel', 2)
        fig = _fig_kinem_corners(heave_x, heave_res, 'rc_height',
                                  'RC Height (mm)', 'Roll Centre Height vs Heave',
                                  corners=('FL', 'RL'))
        _embed_figure(doc, fig, caption='Front and rear RC height. Migration under heave drives jacking force.')
        _analysis_box(doc, _analyze_heave_rc(heave_res, heave_x))
        _rationale_box(doc)

        # 2d — Anti-dive / Anti-squat
        _heading(doc, '2.4 — Anti-Dive & Anti-Squat', 2)
        fig, axes = plt.subplots(1, 2, figsize=(REPORT_FIG_W, 3.0), facecolor='white')
        for ax, key, ttl in zip(axes, ['anti_dive', 'anti_squat'],
                                  ['Anti-Dive (%)', 'Anti-Squat (%)']):
            for lbl in ('FL', 'FR', 'RL', 'RR'):
                d = heave_res.get(lbl, {}).get(key)
                if d is not None and np.any(np.isfinite(d)):
                    ax.plot(heave_x, d, color=_CORNER_COLOR[lbl],
                            linestyle=_CORNER_LS[lbl], linewidth=1.6, label=lbl)
            _ax_style(ax, 'Wheel Travel (mm)', '%', ttl)
            ax.legend(fontsize=8)
        fig.tight_layout(pad=0.5)
        _embed_figure(doc, fig,
                      caption='Anti-dive (front) and anti-squat (rear). 100 % = full geometric resistance to pitch.')
        ad = heave_res.get('FL', {}).get('anti_dive')
        as_ = heave_res.get('RL', {}).get('anti_squat')
        ad_txt = ''
        if ad is not None and np.any(np.isfinite(ad)):
            ad_txt += f'Anti-dive (design pos): {float(ad[len(ad)//2]):.1f} %.  '
        if as_ is not None and np.any(np.isfinite(as_)):
            ad_txt += f'Anti-squat (design pos): {float(as_[len(as_)//2]):.1f} %.'
        _analysis_box(doc, ad_txt.strip() or 'Anti-dive / squat data unavailable.')
        _rationale_box(doc)

        # 2e — Motion Ratio
        _heading(doc, '2.5 — Motion Ratio vs Wheel Travel', 2)
        fig = _fig_kinem_corners(heave_x, heave_res, 'motion_ratio',
                                  'Motion Ratio (—)', 'Motion Ratio vs Heave',
                                  corners=('FL', 'RL'))
        _embed_figure(doc, fig,
                      caption='MR = d(spring) / d(wheel). < 1 = spring moves less than wheel.')
        mr_fl = heave_res.get('FL', {}).get('motion_ratio')
        mr_rl = heave_res.get('RL', {}).get('motion_ratio')
        mr_txt = ''
        for axle_l, mr_d in [('Front', mr_fl), ('Rear', mr_rl)]:
            if mr_d is not None and np.any(np.isfinite(mr_d)):
                mid = len(mr_d) // 2
                mr_txt += f'{axle_l} MR at design pos: {float(mr_d[mid]):.3f}.  '
        _analysis_box(doc, mr_txt.strip() or 'Motion ratio data unavailable.')
        _rationale_box(doc)

    # =========================================================================
    # SECTION 3 — KINEMATICS: ROLL
    # =========================================================================
    if len(roll_x) and roll_res:
        _prog('Kinematics — roll sweep…', 30)
        _heading(doc, '3 — Suspension Kinematics: Roll', 1)
        _body(doc, (
            f'Body roll sweep {roll_x[0]:.1f} ° to {roll_x[-1]:.1f} °. '
            'Shows camber recovery — how well the outer tire stays upright as the body rolls.'
        ), size_pt=10)
        _hr(doc)

        _heading(doc, '3.1 — Camber vs Roll Angle', 2)
        fig = _fig_kinem_corners(roll_x, roll_res, 'camber',
                                  'Camber (°)', 'Camber vs Body Roll',
                                  xlabel='Body Roll Angle (°)')
        _embed_figure(doc, fig,
                      caption='Ideal: outer tire (FL at positive roll) gains negative camber — stays flat on road.')
        # Analysis: check outer tire camber at max roll
        fl_camber = roll_res.get('FL', {}).get('camber')
        if fl_camber is not None and np.any(np.isfinite(fl_camber)):
            max_roll_idx = np.argmax(roll_x) if np.any(roll_x > 0) else -1
            if max_roll_idx >= 0 and max_roll_idx < len(fl_camber):
                fl_c = float(fl_camber[max_roll_idx])
                roll_at = float(roll_x[max_roll_idx])
                if fl_c < -0.5:
                    roll_txt = (f'FL (outside) camber: {fl_c:.2f} ° at {roll_at:.1f} ° body roll — '
                                f'good negative recovery.')
                elif fl_c < 0:
                    roll_txt = (f'FL camber: {fl_c:.2f} ° at {roll_at:.1f} ° body roll — '
                                f'slight negative gain. Consider geometry tuning.')
                else:
                    roll_txt = (f'FL camber: {fl_c:.2f} ° at {roll_at:.1f} ° body roll — '
                                f'going positive on the outer tire. Review UCA geometry.')
            else:
                roll_txt = 'Insufficient roll range data.'
        else:
            roll_txt = 'Camber data unavailable for roll sweep.'
        _analysis_box(doc, roll_txt)
        _rationale_box(doc)

        _heading(doc, '3.2 — RC Height vs Roll Angle', 2)
        fig = _fig_kinem_corners(roll_x, roll_res, 'rc_height',
                                  'RC Height (mm)', 'RC Height vs Body Roll',
                                  xlabel='Body Roll Angle (°)',
                                  corners=('FL', 'RL'))
        _embed_figure(doc, fig, caption='RC migration during body roll couples with jacking force.')
        _analysis_box(doc, _analyze_heave_rc(roll_res, roll_x))
        _rationale_box(doc)

    # =========================================================================
    # SECTION 4 — STEADY-STATE CORNERING
    # =========================================================================
    _prog('Steady-state cornering sweep…', 45)
    _heading(doc, '4 — Steady-State Cornering', 1)
    g_corn = np.asarray(dyn_corn.get('lateral_g', []))
    _body(doc, (
        f'Lateral-g sweep 0 – {g_corn[-1]:.1f} g (pure cornering, lon-g = 0). '
        'Steady-state load transfer, roll, utilization.'
        if len(g_corn) else 'Cornering data not available.'
    ), size_pt=10)
    _hr(doc)

    if len(g_corn) and dyn_corn:
        corn_analysis = _analyze_cornering(dyn_corn)

        # 4a — Roll Angle
        _heading(doc, '4.1 — Body Roll Angle', 2)
        roll_arr = np.asarray(dyn_corn.get('roll_angle_deg', np.zeros_like(g_corn)))
        fig = _fig_dyn_multi(g_corn,
                              [('Roll', roll_arr, '#1565C0', '-')],
                              'Lateral g', 'Body Roll Angle vs Lateral g')
        fig.axes[0].set_ylabel('Roll Angle (°)', fontsize=9)
        _embed_figure(doc, fig)
        _analysis_box(doc, corn_analysis['roll'])
        _rationale_box(doc)

        # 4b — Tire Loads
        _heading(doc, '4.2 — Corner Tire Loads (Fz)', 2)
        fz_series = [(c, np.asarray(dyn_corn.get(f'Fz_{c}', np.zeros_like(g_corn))),
                      _CORNER_COLOR[c], _CORNER_LS[c]) for c in ('FL', 'FR', 'RL', 'RR')]
        fig = _fig_dyn_multi(g_corn, fz_series, 'Lateral g', 'Corner Loads vs Lateral g')
        fig.axes[0].set_ylabel('Fz (N)', fontsize=9)
        _embed_figure(doc, fig, caption='Inner tires unload; outer tires load up. LT ratio determines handling balance.')
        _analysis_box(doc, corn_analysis['lt'])
        _rationale_box(doc)

        # 4c — Tire Utilization
        _heading(doc, '4.3 — Tire Utilization (Friction Circle)', 2)
        util_series = [(c, np.asarray(dyn_corn.get(f'utilization_{c}', np.zeros_like(g_corn))),
                        _CORNER_COLOR[c], _CORNER_LS[c]) for c in ('FL', 'FR', 'RL', 'RR')]
        fig = _fig_dyn_multi(g_corn, util_series, 'Lateral g', 'Tire Utilization vs Lateral g')
        fig.axes[0].set_ylabel('Utilization (—)', fontsize=9)
        fig.axes[0].axhline(1.0, color='#C62828', linewidth=1.0, linestyle='--', alpha=0.7)
        fig.axes[0].set_ylim(0, 1.15)
        _embed_figure(doc, fig, caption='1.0 = tire fully saturated. Dashed line = limit.')
        _analysis_box(doc, corn_analysis['util'])
        _rationale_box(doc)

        # 4d — Load Transfer Breakdown
        _heading(doc, '4.4 — Lateral Load Transfer Breakdown', 2)
        lt_series = [
            ('Elastic Front',   np.asarray(dyn_corn.get('elastic_lt_front_N',   np.zeros_like(g_corn))), _ELF, '-'),
            ('Elastic Rear',    np.asarray(dyn_corn.get('elastic_lt_rear_N',    np.zeros_like(g_corn))), _ELR, '--'),
            ('Geometric Front', np.asarray(dyn_corn.get('geometric_lt_front_N', np.zeros_like(g_corn))), _GLF, '-.'),
            ('Geometric Rear',  np.asarray(dyn_corn.get('geometric_lt_rear_N',  np.zeros_like(g_corn))), _GLR, ':'),
        ]
        fig = _fig_dyn_multi(g_corn, lt_series, 'Lateral g', 'Load Transfer Breakdown')
        fig.axes[0].set_ylabel('LT (N)', fontsize=9)
        _embed_figure(doc, fig,
                      caption='Elastic LT via springs + ARBs. Geometric LT via RC height. '
                               'Higher geometric LT front = understeer tendency.')
        # LLTD analysis
        el_f = np.asarray(dyn_corn.get('elastic_lt_front_N',   np.zeros_like(g_corn)))
        el_r = np.asarray(dyn_corn.get('elastic_lt_rear_N',    np.zeros_like(g_corn)))
        ge_f = np.asarray(dyn_corn.get('geometric_lt_front_N', np.zeros_like(g_corn)))
        ge_r = np.asarray(dyn_corn.get('geometric_lt_rear_N',  np.zeros_like(g_corn)))
        idx_1g = int(np.argmin(np.abs(g_corn - 1.0)))
        total_f = float(el_f[idx_1g] + ge_f[idx_1g])
        total_r = float(el_r[idx_1g] + ge_r[idx_1g])
        total   = total_f + total_r
        lltd_f  = _pct(total_f, total)
        lt_txt  = (f'At 1 g: LLTD front {lltd_f:.0f} % / rear {100-lltd_f:.0f} %. '
                   f'{"Front-biased → understeer tendency" if lltd_f > 55 else "Rear-biased → oversteer tendency" if lltd_f < 45 else "Balanced split."}')
        _analysis_box(doc, lt_txt)
        _rationale_box(doc)

        # 4e — Pitch during cornering
        _heading(doc, '4.5 — Pitch Angle during Cornering', 2)
        pitch_arr = np.asarray(dyn_corn.get('pitch_angle_deg', np.zeros_like(g_corn)))
        fig = _fig_dyn_multi(g_corn, [('Pitch', pitch_arr, '#6A1B9A', '-')],
                              'Lateral g', 'Pitch Angle vs Lateral g')
        fig.axes[0].set_ylabel('Pitch (°)', fontsize=9)
        _embed_figure(doc, fig, caption='Small pitch change during pure cornering from Fz asymmetry.')
        _analysis_box(doc, corn_analysis['pitch'])
        _rationale_box(doc)

    # =========================================================================
    # SECTION 5 — ACCELERATION TRAJECTORY
    # =========================================================================
    _prog('Acceleration trajectory…', 65)
    _heading(doc, '5 — Straight-Line Acceleration', 1)
    t_accel = np.asarray(dyn_accel.get('time_s', []))
    v_accel = np.asarray(dyn_accel.get('speed_mph', []))
    _body(doc, (
        f'Full-throttle trajectory from 0 mph. Duration: {float(t_accel[-1]):.1f} s, '
        f'terminal: {float(v_accel[-1]):.0f} mph.'
        if len(t_accel) > 1 else 'Acceleration data not available.'
    ), size_pt=10)
    _hr(doc)

    if len(t_accel) > 1:
        # 5a — Speed + g on two axes
        _heading(doc, '5.1 — Speed & Achieved Longitudinal g', 2)
        g_arr = np.asarray(dyn_accel.get('longitudinal_g', np.zeros_like(t_accel)))
        fig = _fig_two_panel(
            t_accel,
            [('Speed', v_accel, _FL, '-')], [('Lon-g', g_arr, _RL, '-')],
            'Time (s)', 'Speed (mph)', 'Lon-g (g)', 'Acceleration Trajectory')
        _embed_figure(doc, fig)
        _analysis_box(doc, _analyze_traj(dyn_accel, 'Acceleration'))
        _rationale_box(doc)

        # 5b — Pitch + Corner loads
        _heading(doc, '5.2 — Pitch & Tire Loads', 2)
        pitch_a = np.asarray(dyn_accel.get('pitch_angle_deg', np.zeros_like(t_accel)))
        fz_top = [('Pitch', pitch_a, '#6A1B9A', '-')]
        fz_bot = [(c, np.asarray(dyn_accel.get(f'Fz_{c}', np.zeros_like(t_accel))),
                   _CORNER_COLOR[c], _CORNER_LS[c]) for c in ('FL', 'FR', 'RL', 'RR')]
        fig = _fig_two_panel(t_accel, fz_top, fz_bot,
                              'Time (s)', 'Pitch (°)', 'Fz (N)',
                              'Pitch & Corner Loads — Acceleration')
        _embed_figure(doc, fig, caption='Nose rises as weight shifts to driven rear axle.')
        fz_fl_0 = float(np.asarray(dyn_accel.get('Fz_FL', [0]))[0])
        fz_rl_0 = float(np.asarray(dyn_accel.get('Fz_RL', [0]))[0])
        pitch_0 = float(pitch_a[0]) if len(pitch_a) else 0
        _analysis_box(doc,
            f'At launch: front Fz {fz_fl_0:.0f} N, rear Fz {fz_rl_0:.0f} N. '
            f'Pitch {pitch_0:.2f} ° — driven by weight transfer to rear.')
        _rationale_box(doc)

    # =========================================================================
    # SECTION 6 — BRAKING
    # =========================================================================
    _prog('Braking trajectory…', 80)
    _heading(doc, '6 — Straight-Line Braking', 1)
    t_brake = np.asarray(dyn_brake.get('time_s', []))
    v_brake = np.asarray(dyn_brake.get('speed_mph', []))
    _body(doc, (
        f'Max-braking trajectory from {float(v_brake[0]):.0f} mph. '
        f'Stops in {float(t_brake[-1]):.2f} s.'
        if len(t_brake) > 1 else 'Braking data not available.'
    ), size_pt=10)
    _hr(doc)

    if len(t_brake) > 1:
        _heading(doc, '6.1 — Speed & Decel g', 2)
        g_b = np.asarray(dyn_brake.get('longitudinal_g', np.zeros_like(t_brake)))
        fig = _fig_two_panel(
            t_brake,
            [('Speed', v_brake, _FL, '-')], [('Lon-g', g_b, _RL, '-')],
            'Time (s)', 'Speed (mph)', 'Lon-g (g)', 'Braking Trajectory')
        _embed_figure(doc, fig)
        _analysis_box(doc, _analyze_traj(dyn_brake, 'Braking'))
        _rationale_box(doc)

        _heading(doc, '6.2 — Pitch & Tire Loads', 2)
        pitch_b = np.asarray(dyn_brake.get('pitch_angle_deg', np.zeros_like(t_brake)))
        fz_top2 = [('Pitch', pitch_b, '#6A1B9A', '-')]
        fz_bot2 = [(c, np.asarray(dyn_brake.get(f'Fz_{c}', np.zeros_like(t_brake))),
                    _CORNER_COLOR[c], _CORNER_LS[c]) for c in ('FL', 'FR', 'RL', 'RR')]
        fig = _fig_two_panel(t_brake, fz_top2, fz_bot2,
                              'Time (s)', 'Pitch (°)', 'Fz (N)',
                              'Pitch & Corner Loads — Braking')
        _embed_figure(doc, fig, caption='Nose dives as weight shifts forward. Rear tires unload.')
        _analysis_box(doc,
            f'Front Fz peaks at {float(np.nanmax(np.asarray(dyn_brake.get("Fz_FL", [0])))):.0f} N (FL). '
            f'Rear unloads to {float(np.nanmin(np.asarray(dyn_brake.get("Fz_RL", [0])))):.0f} N min (RL).')
        _rationale_box(doc)

    # =========================================================================
    # SECTION 7 — COMPONENT LOADS
    # =========================================================================
    loads_data = data.get('loads')
    if loads_data and loads_data.get('corners'):
        _prog('Component loads…', 90)
        _heading(doc, '7 — Component Loads', 1)
        lat_g_l = loads_data.get('lat_g', 0.0)
        lon_g_l = loads_data.get('lon_g', 0.0)
        _body(doc, (
            f'Member forces and joint reactions at {lat_g_l:.2f} g lateral, '
            f'{lon_g_l:.2f} g longitudinal (current panel state).'
        ), size_pt=10)
        _hr(doc)

        corners_loads = loads_data['corners']

        # 7a — Axial member forces table
        _heading(doc, '7.1 — Suspension Member Axial Forces', 2)
        _body(doc, 'Positive = tension (pulled apart). Negative = compression.', size_pt=9.5, italic=True)

        axial_attrs = [
            ('UCA front arm',  'uca_front_N'),
            ('UCA rear arm',   'uca_rear_N'),
            ('LCA front arm',  'lca_front_N'),
            ('LCA rear arm',   'lca_rear_N'),
            ('Tie rod',        'tierod_N'),
            ('Pushrod',        'pushrod_N'),
            ('Spring (comp+)', 'spring_force_N'),
        ]
        # Build a 5-column table: Member | FL | FR | RL | RR
        corner_labels = ['FL', 'FR', 'RL', 'RR']
        tbl = doc.add_table(rows=len(axial_attrs) + 1, cols=5)
        tbl.style = 'Table Grid'
        # Header
        for ci, hdr in enumerate(['Member', 'FL', 'FR', 'RL', 'RR']):
            c = tbl.cell(0, ci)
            c.text = hdr
            _set_cell_bg(c, '#1565C0')
            c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            c.paragraphs[0].runs[0].font.bold = True
            c.paragraphs[0].runs[0].font.name = 'Arial'
            c.paragraphs[0].runs[0].font.size = Pt(9)
        # Data rows
        for r_idx, (label, attr) in enumerate(axial_attrs):
            row_cells = tbl.rows[r_idx + 1].cells
            bg = '#FFFFFF' if r_idx % 2 == 0 else '#F5F5F5'
            row_cells[0].text = label
            _set_cell_bg(row_cells[0], bg)
            row_cells[0].paragraphs[0].runs[0].font.name = 'Arial'
            row_cells[0].paragraphs[0].runs[0].font.size = Pt(9)
            for ci, clbl in enumerate(corner_labels, start=1):
                cl = corners_loads.get(clbl)
                val = getattr(cl, attr, 0.0) if cl else 0.0
                row_cells[ci].text = f'{val:+.0f}'
                _set_cell_bg(row_cells[ci], bg)
                row_cells[ci].paragraphs[0].runs[0].font.name = 'Arial'
                row_cells[ci].paragraphs[0].runs[0].font.size = Pt(9)
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

        # 7b — Ball joint reactions table
        _heading(doc, '7.2 — Ball Joint Reactions (V=up+, H=fwd+)', 2)
        bj_attrs = [
            ('UCA BJ V',     'uca_bj_V'),
            ('UCA BJ H',     'uca_bj_H'),
            ('LCA BJ V',     'lca_bj_V'),
            ('LCA BJ H',     'lca_bj_H'),
            ('Tie Rod BJ V', 'tierod_bj_V'),
            ('Tie Rod BJ H', 'tierod_bj_H'),
            ('Pushrod BJ V', 'pushrod_bj_V'),
            ('Pushrod BJ H', 'pushrod_bj_H'),
        ]
        tbl2 = doc.add_table(rows=len(bj_attrs) + 1, cols=5)
        tbl2.style = 'Table Grid'
        for ci, hdr in enumerate(['Joint', 'FL', 'FR', 'RL', 'RR']):
            c = tbl2.cell(0, ci)
            c.text = hdr
            _set_cell_bg(c, '#1565C0')
            c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            c.paragraphs[0].runs[0].font.bold = True
            c.paragraphs[0].runs[0].font.name = 'Arial'
            c.paragraphs[0].runs[0].font.size = Pt(9)
        for r_idx, (label, attr) in enumerate(bj_attrs):
            row_cells = tbl2.rows[r_idx + 1].cells
            bg = '#FFFFFF' if r_idx % 2 == 0 else '#F5F5F5'
            row_cells[0].text = label
            _set_cell_bg(row_cells[0], bg)
            row_cells[0].paragraphs[0].runs[0].font.name = 'Arial'
            row_cells[0].paragraphs[0].runs[0].font.size = Pt(9)
            for ci, clbl in enumerate(corner_labels, start=1):
                cl = corners_loads.get(clbl)
                val = getattr(cl, attr, 0.0) if cl else 0.0
                row_cells[ci].text = f'{val:+.0f}'
                _set_cell_bg(row_cells[ci], bg)
                row_cells[ci].paragraphs[0].runs[0].font.name = 'Arial'
                row_cells[ci].paragraphs[0].runs[0].font.size = Pt(9)
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

        # 7c — Bearing & Caliper
        _heading(doc, '7.3 — Bearing & Caliper Bolt Loads', 2)
        bc_attrs = [
            ('Bearing inner V', 'bearing_inner_V'),
            ('Bearing inner H', 'bearing_inner_H'),
            ('Bearing outer V', 'bearing_outer_V'),
            ('Bearing outer H', 'bearing_outer_H'),
            ('Caliper upper V', 'caliper_upper_V'),
            ('Caliper upper H', 'caliper_upper_H'),
            ('Caliper lower V', 'caliper_lower_V'),
            ('Caliper lower H', 'caliper_lower_H'),
        ]
        tbl3 = doc.add_table(rows=len(bc_attrs) + 1, cols=5)
        tbl3.style = 'Table Grid'
        for ci, hdr in enumerate(['Load', 'FL', 'FR', 'RL', 'RR']):
            c = tbl3.cell(0, ci)
            c.text = hdr
            _set_cell_bg(c, '#1565C0')
            c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            c.paragraphs[0].runs[0].font.bold = True
            c.paragraphs[0].runs[0].font.name = 'Arial'
            c.paragraphs[0].runs[0].font.size = Pt(9)
        for r_idx, (label, attr) in enumerate(bc_attrs):
            row_cells = tbl3.rows[r_idx + 1].cells
            bg = '#FFFFFF' if r_idx % 2 == 0 else '#F5F5F5'
            row_cells[0].text = label
            _set_cell_bg(row_cells[0], bg)
            row_cells[0].paragraphs[0].runs[0].font.name = 'Arial'
            row_cells[0].paragraphs[0].runs[0].font.size = Pt(9)
            for ci, clbl in enumerate(corner_labels, start=1):
                cl = corners_loads.get(clbl)
                val = getattr(cl, attr, 0.0) if cl else 0.0
                row_cells[ci].text = f'{val:+.0f}'
                _set_cell_bg(row_cells[ci], bg)
                row_cells[ci].paragraphs[0].runs[0].font.name = 'Arial'
                row_cells[ci].paragraphs[0].runs[0].font.size = Pt(9)
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

        # Analysis: peak forces summary
        peak_push = 0.0
        peak_tierod = 0.0
        peak_spring = 0.0
        for clbl in corner_labels:
            cl = corners_loads.get(clbl)
            if cl:
                peak_push = max(peak_push, abs(getattr(cl, 'pushrod_N', 0)))
                peak_tierod = max(peak_tierod, abs(getattr(cl, 'tierod_N', 0)))
                peak_spring = max(peak_spring, abs(getattr(cl, 'spring_force_N', 0)))
        _analysis_box(doc,
            f'Peak pushrod: {peak_push:.0f} N.  '
            f'Peak tie rod: {peak_tierod:.0f} N.  '
            f'Peak spring: {peak_spring:.0f} N.  '
            f'At {lat_g_l:.2f}g lat / {lon_g_l:.2f}g lon.')
        _rationale_box(doc)

    # =========================================================================
    # SAVE
    # =========================================================================
    _prog('Saving…', 96)
    doc.save(output_path)
    _prog('Done.', 100)
